from flask import Flask, render_template, request, jsonify
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.tokenize import sent_tokenize
import io
import os
import traceback
import requests as req_lib
from concurrent.futures import ThreadPoolExecutor, as_completed

# Download NLTK data to /tmp so it works on Vercel (read-only filesystem except /tmp)
_NLTK_DIR = '/tmp/nltk_data'
os.makedirs(_NLTK_DIR, exist_ok=True)
nltk.data.path.insert(0, _NLTK_DIR)
nltk.download('punkt',     download_dir=_NLTK_DIR, quiet=True)
nltk.download('punkt_tab', download_dir=_NLTK_DIR, quiet=True)

app = Flask(__name__)

# ── File-type detection & text extraction ───────────────────────────────────────

MIME_TO_EXT = {
    'application/pdf': '.pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
    'application/msword': '.doc',
    'text/plain': '.txt',
}


def detect_ext(file):
    fname = os.path.basename(file.filename or '').lower()
    ext = os.path.splitext(fname)[1]
    if ext in ('.txt', '.pdf', '.docx', '.doc'):
        return ext
    mime = getattr(file, 'content_type', None) or getattr(file, 'mimetype', None) or ''
    return MIME_TO_EXT.get(mime.split(';')[0].strip(), ext or '.unknown')


def read_file_bytes(file):
    stream = getattr(file, 'stream', None)
    if stream is not None:
        stream.seek(0)
        return stream.read()
    return file.read()


def extract_text_from_file(file):
    ext = detect_ext(file)
    data = read_file_bytes(file)
    if not data:
        raise ValueError('The uploaded file appears to be empty.')

    if ext == '.txt':
        try:
            return data.decode('utf-8')
        except UnicodeDecodeError:
            return data.decode('latin-1', errors='replace')

    elif ext == '.pdf':
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        result = '\n'.join(text_parts).strip()
        if not result:
            raise ValueError(
                'No text could be extracted from this PDF. '
                'It may be a scanned/image-only PDF — try the "Paste Text" tab instead.'
            )
        return result

    elif ext == '.docx':
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return '\n'.join(parts)

    elif ext == '.doc':
        raise ValueError(
            'Old .doc format is not supported. '
            'Please open in Word and save as .docx, then upload again.'
        )

    else:
        raise ValueError(
            f'Unsupported file type "{ext}". '
            'Please upload a .txt, .pdf, or .docx file.'
        )


# ── Core similarity functions ───────────────────────────────────────────────────

def calculate_similarity(text1, text2):
    documents = [text1, text2]
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(documents)
    return float(cosine_similarity(tfidf_matrix)[0][1])


def sentence_level_similarity(text1, text2, threshold=0.5):
    sentences1 = sent_tokenize(text1)
    sentences2 = sent_tokenize(text2)
    pairs = []
    for s1 in sentences1:
        best_score, best_s2 = 0, ''
        for s2 in sentences2:
            if len(s1.split()) < 3 or len(s2.split()) < 3:
                continue
            try:
                tfidf = TfidfVectorizer(stop_words='english').fit_transform([s1, s2])
                score = cosine_similarity(tfidf)[0][1]
                if score > best_score:
                    best_score, best_s2 = score, s2
            except Exception:
                continue
        if best_score > threshold:
            pairs.append({'s1': s1, 's2': best_s2, 'score': round(best_score * 100, 1)})
    pairs.sort(key=lambda x: x['score'], reverse=True)
    return pairs[:20]


# ── Online check helpers ────────────────────────────────────────────────────────

def pick_search_queries(text, n=5):
    """Pick the n most distinctive sentences from the text to use as search queries."""
    sentences = [s for s in sent_tokenize(text) if len(s.split()) >= 8]
    if not sentences:
        return [text[:200]]

    if len(sentences) <= n:
        return sentences

    try:
        # Score sentences by average TF-IDF weight of their words
        vectorizer = TfidfVectorizer(stop_words='english')
        tfidf = vectorizer.fit_transform(sentences)
        scores = tfidf.mean(axis=1).A1
        top_indices = scores.argsort()[::-1][:n]
        return [sentences[i] for i in sorted(top_indices)]
    except Exception:
        # Fallback: evenly spaced sentences
        step = len(sentences) // n
        return [sentences[i * step] for i in range(n)]


def search_web(queries, max_results_per_query=3):
    """Search DuckDuckGo for each query and return a deduplicated list of URLs."""
    from duckduckgo_search import DDGS
    urls = []
    seen = set()
    try:
        with DDGS() as ddgs:
            for query in queries:
                try:
                    results = list(ddgs.text(
                        f'"{query[:120]}"',
                        max_results=max_results_per_query
                    ))
                    for r in results:
                        url = r.get('href') or r.get('url', '')
                        if url and url not in seen:
                            seen.add(url)
                            urls.append(url)
                except Exception:
                    continue
    except Exception:
        pass
    return urls[:15]  # Cap at 15 total URLs


def fetch_page_text(url, timeout=6):
    """Fetch a URL and return clean text. Returns None on any failure."""
    try:
        headers = {
            'User-Agent': (
                'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                'AppleWebKit/537.36 (KHTML, like Gecko) '
                'Chrome/120.0.0.0 Safari/537.36'
            )
        }
        resp = req_lib.get(url, headers=headers, timeout=timeout, allow_redirects=True)
        if resp.status_code != 200:
            return None

        content_type = resp.headers.get('Content-Type', '')
        if 'text/html' not in content_type and 'text/plain' not in content_type:
            return None

        from bs4 import BeautifulSoup
        soup = BeautifulSoup(resp.text, 'lxml')

        # Remove navigation, ads, scripts, styles
        for tag in soup(['script', 'style', 'nav', 'header', 'footer',
                         'aside', 'form', 'noscript', 'iframe']):
            tag.decompose()

        # Try to get the main content first
        main = soup.find('main') or soup.find('article') or soup.find('body')
        raw = main.get_text(separator=' ', strip=True) if main else soup.get_text()

        # Clean up excessive whitespace
        lines = [line.strip() for line in raw.splitlines() if line.strip()]
        text = ' '.join(lines)
        return text if len(text.split()) >= 30 else None

    except Exception:
        return None


def check_one_source(args):
    """Thread worker: fetch a URL and compute similarity against the document."""
    url, doc_text = args
    page_text = fetch_page_text(url)
    if not page_text:
        return None
    try:
        score = calculate_similarity(doc_text, page_text)
        if score < 0.05:
            return None
        pairs = sentence_level_similarity(doc_text, page_text, threshold=0.55)
        label_map = [
            (0.8, 'High Plagiarism',    'high'),
            (0.5, 'Moderate Similarity','medium'),
            (0.2, 'Low Similarity',     'low'),
            (0.0, 'Minimal Similarity', 'minimal'),
        ]
        label, color = next(
            (lbl, clr) for thresh, lbl, clr in label_map if score >= thresh
        )
        return {
            'url': url,
            'similarity': round(score * 100, 1),
            'label': label,
            'color': color,
            'pairs': pairs[:5],   # Keep top 5 sentence matches per source
        }
    except Exception:
        return None


# ── Routes ──────────────────────────────────────────────────────────────────────

@app.route('/')
def home():
    return render_template('index.html')


@app.route('/check', methods=['POST'])
def check():
    try:
        text1, text2 = '', ''
        mode = request.form.get('mode', 'file')

        if mode == 'text':
            text1 = request.form.get('text1', '').strip()
            text2 = request.form.get('text2', '').strip()
        else:
            file1 = request.files.get('file1')
            file2 = request.files.get('file2')
            if not file1 or not file2:
                return jsonify({'error': 'Please upload both files.'}), 400
            try:
                text1 = extract_text_from_file(file1).strip()
            except ValueError as e:
                return jsonify({'error': f'Document 1: {e}'}), 400
            except Exception as e:
                traceback.print_exc()
                return jsonify({'error': f'Document 1 could not be read: {e}'}), 400
            try:
                text2 = extract_text_from_file(file2).strip()
            except ValueError as e:
                return jsonify({'error': f'Document 2: {e}'}), 400
            except Exception as e:
                traceback.print_exc()
                return jsonify({'error': f'Document 2 could not be read: {e}'}), 400

        if not text1 or not text2:
            return jsonify({'error': 'Both inputs must have content.'}), 400
        if len(text1.split()) < 5 or len(text2.split()) < 5:
            return jsonify({'error': 'Each document must have at least 5 words.'}), 400

        score = calculate_similarity(text1, text2)
        similarity = round(score * 100, 1)
        if score > 0.8:   label, color = 'High Plagiarism',    'high'
        elif score > 0.5: label, color = 'Moderate Similarity','medium'
        elif score > 0.2: label, color = 'Low Similarity',     'low'
        else:             label, color = 'Minimal Similarity', 'minimal'

        pairs = sentence_level_similarity(text1, text2)
        return jsonify({
            'similarity': similarity, 'label': label, 'color': color,
            'pairs': pairs, 'word_count_1': len(text1.split()),
            'word_count_2': len(text2.split()),
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': f'Unexpected error: {e}'}), 500


@app.route('/check-online', methods=['POST'])
def check_online():
    try:
        mode = request.form.get('mode', 'file')

        if mode == 'text':
            doc_text = request.form.get('text1', '').strip()
        else:
            file1 = request.files.get('file1')
            if not file1:
                return jsonify({'error': 'Please upload a file.'}), 400
            try:
                doc_text = extract_text_from_file(file1).strip()
            except ValueError as e:
                return jsonify({'error': str(e)}), 400
            except Exception as e:
                traceback.print_exc()
                return jsonify({'error': f'Could not read file: {e}'}), 400

        if not doc_text:
            return jsonify({'error': 'The document appears to be empty.'}), 400
        if len(doc_text.split()) < 20:
            return jsonify({'error': 'Document needs at least 20 words for web search.'}), 400

        # Step 1: pick distinctive queries
        queries = pick_search_queries(doc_text, n=5)

        # Step 2: search the web
        urls = search_web(queries, max_results_per_query=3)
        if not urls:
            return jsonify({'error': 'No web results found. Try again or check your connection.'}), 400

        # Step 3: fetch & compare in parallel (max 8 workers)
        results = []
        with ThreadPoolExecutor(max_workers=8) as ex:
            futures = {ex.submit(check_one_source, (url, doc_text)): url for url in urls}
            for future in as_completed(futures):
                result = future.result()
                if result:
                    results.append(result)

        # Sort by similarity
        results.sort(key=lambda x: x['similarity'], reverse=True)

        return jsonify({
            'sources': results[:10],
            'total_searched': len(urls),
            'total_matched': len(results),
            'word_count': len(doc_text.split()),
            'queries_used': queries,
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': f'Online check failed: {e}'}), 500


if __name__ == '__main__':
    app.run(debug=True)